from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def extract_resume_text(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(content)
    if suffix == ".docx":
        return _extract_docx(content)
    raise ValueError("Unsupported resume format. Upload a PDF or DOCX file.")


def _extract_pdf(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def _extract_docx(content: bytes) -> str:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            table_text.extend(cell.text for cell in row.cells)
    return "\n".join(paragraphs + table_text)
